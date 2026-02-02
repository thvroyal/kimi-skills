#!/usr/bin/env python3
"""Generate Faker Career Analysis DOCX document"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Color scheme - Esports/Gaming theme
PRIMARY_COLOR = RGBColor(0x1a, 0x1a, 0x3e)  # Dark blue
ACCENT_COLOR = RGBColor(0xc8, 0xaa, 0x64)   # Gold
SECONDARY_COLOR = RGBColor(0x2d, 0x1b, 0x4e) # Purple
TEXT_COLOR = RGBColor(0x33, 0x33, 0x33)

def set_cell_shading(cell, color):
    """Set cell background color"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_styled_heading(doc, text, level=1):
    """Add a styled heading"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        if level == 1:
            run.font.color.rgb = PRIMARY_COLOR
            run.font.size = Pt(24)
        elif level == 2:
            run.font.color.rgb = SECONDARY_COLOR
            run.font.size = Pt(16)
        elif level == 3:
            run.font.color.rgb = TEXT_COLOR
            run.font.size = Pt(13)
    return heading

def add_quote_box(doc, quote, author):
    """Add a styled quote box"""
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Cm(1)
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(12)
    
    run = para.add_run(f'"{quote}"')
    run.italic = True
    run.font.color.rgb = TEXT_COLOR
    
    author_para = doc.add_paragraph()
    author_para.paragraph_format.left_indent = Cm(1)
    author_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    author_run = author_para.add_run(f"— {author}")
    author_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    author_run.bold = True

def add_stats_table(doc, stats):
    """Add a statistics table"""
    rows = (len(stats) + 2) // 3
    table = doc.add_table(rows=rows, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    idx = 0
    for i in range(rows):
        for j in range(3):
            if idx < len(stats):
                cell = table.cell(i, j)
                label, value = stats[idx]
                
                # Value
                p1 = cell.paragraphs[0]
                p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run1 = p1.add_run(str(value))
                run1.bold = True
                run1.font.size = Pt(18)
                run1.font.color.rgb = ACCENT_COLOR
                
                # Label
                p2 = cell.add_paragraph()
                p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run2 = p2.add_run(label)
                run2.font.size = Pt(9)
                run2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                
                idx += 1
    
    doc.add_paragraph()
    return table

def add_data_table(doc, headers, data, highlight_first=False):
    """Add a data table with headers"""
    table = doc.add_table(rows=1 + len(data), cols=len(headers))
    table.style = 'Table Grid'
    
    # Header row
    header_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = header
        set_cell_shading(cell, 'f0f0f0')
        cell.paragraphs[0].runs[0].bold = True
    
    # Data rows
    for row_idx, row_data in enumerate(data):
        row = table.rows[row_idx + 1]
        for col_idx, value in enumerate(row_data):
            cell = row.cells[col_idx]
            cell.text = str(value)
            if highlight_first and row_idx == 0:
                set_cell_shading(cell, 'fff8e7')
                cell.paragraphs[0].runs[0].bold = True
    
    doc.add_paragraph()
    return table

def create_faker_document():
    """Create the complete Faker career analysis document"""
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # =========================================================================
    # COVER PAGE
    # =========================================================================
    # Add spacing at top
    for _ in range(8):
        doc.add_paragraph()
    
    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("LEAGUE OF LEGENDS")
    run.font.size = Pt(14)
    run.font.color.rgb = ACCENT_COLOR
    run.font.bold = True
    
    # Main title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("FAKER")
    run.font.size = Pt(56)
    run.font.color.rgb = PRIMARY_COLOR
    run.bold = True
    
    # Tagline
    tagline = doc.add_paragraph()
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = tagline.add_run("The Unkillable Demon King")
    run.font.size = Pt(18)
    run.italic = True
    run.font.color.rgb = TEXT_COLOR
    
    # Spacing
    for _ in range(2):
        doc.add_paragraph()
    
    # Subtitle info
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run("Complete Career Analysis & Visual Infographic\n2013 – 2025 | 12+ Years of Dominance")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    
    # Key stats on cover
    for _ in range(3):
        doc.add_paragraph()
    
    stats_para = doc.add_paragraph()
    stats_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    stats_text = "6 World Titles  •  10 LCK Titles  •  1600+ Games Played"
    run = stats_para.add_run(stats_text)
    run.font.size = Pt(14)
    run.font.color.rgb = ACCENT_COLOR
    run.bold = True
    
    doc.add_page_break()

    # =========================================================================
    # TABLE OF CONTENTS
    # =========================================================================
    add_styled_heading(doc, "Contents", 1)
    
    toc_items = [
        ("1. Introduction: The GOAT of Esports", 1),
        ("   1.1 Player Profile", 2),
        ("   1.2 Career Overview", 2),
        ("2. The Rise: 2013-2014", 1),
        ("   2.1 Rookie Sensation", 2),
        ("   2.2 First World Championship", 2),
        ("3. Dynasty Era: 2015-2017", 1),
        ("   3.1 Back-to-Back World Titles", 2),
        ("   3.2 The 2017 Heartbreak", 2),
        ("4. The Dark Age: 2018-2022", 1),
        ("   4.1 Rebuilding Years", 2),
        ("   4.2 Near Misses", 2),
        ("5. The Return of the King: 2023-2025", 1),
        ("   5.1 Fourth World Title", 2),
        ("   5.2 Historic Three-Peat", 2),
        ("6. Statistical Analysis", 1),
        ("   6.1 Career Statistics", 2),
        ("   6.2 Champion Pool", 2),
        ("7. Legacy & Impact", 1),
        ("8. Conclusion", 1),
    ]
    
    for item, level in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(item)
        if level == 1:
            run.bold = True
        run.font.size = Pt(11 if level == 1 else 10)
    
    hint = doc.add_paragraph()
    hint.paragraph_format.space_before = Pt(20)
    run = hint.add_run("(Hint: Use Ctrl+Click on headings to navigate to sections)")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.italic = True
    
    doc.add_page_break()

    # =========================================================================
    # SECTION 1: INTRODUCTION
    # =========================================================================
    add_styled_heading(doc, "1. Introduction: The GOAT of Esports", 1)
    
    doc.add_paragraph(
        'In the annals of competitive gaming, no name resonates with the same reverence and awe as '
        'Lee "Faker" Sang-hyeok. Often referred to as the "Michael Jordan of esports" and the '
        '"Unkillable Demon King," Faker has transcended the boundaries of professional League of '
        'Legends to become a cultural phenomenon and the undisputed greatest player in the history of the game.'
    )
    
    add_quote_box(doc, 
        "Faker is truly the most iconic and influential figure in League of Legends and esports history. "
        "His name is synonymous with excellence in LoL Esports.",
        "John Needham, President of Esports at Riot Games"
    )
    
    add_styled_heading(doc, "1.1 Player Profile", 2)
    
    add_stats_table(doc, [
        ("Korean Name", "이상혁"),
        ("Date of Birth", "May 7, 1996"),
        ("Birthplace", "Seoul"),
        ("Position", "Mid Laner"),
        ("Team", "T1 (2013-Present)"),
        ("Career Length", "12+ Years"),
    ])
    
    doc.add_paragraph(
        "Born on May 7, 1996, in Seoul, South Korea, Faker joined SK Telecom T1 (now T1) in 2013 "
        "and has remained with the organization for his entire professional career—the longest "
        "tenure of any player with a single organization in esports history."
    )
    
    add_styled_heading(doc, "1.2 Career Overview", 2)
    
    # Achievement badges as bullet points
    achievements = [
        "6× World Champion (2013, 2015, 2016, 2023, 2024, 2025)",
        "10× LCK Champion",
        "2× MSI Champion (2016, 2017)",
        "Hall of Legends Inaugural Inductee (2024)",
        "3× Best Esports Athlete - The Game Awards (2017, 2023, 2024)",
        "2× Worlds Finals MVP (2016, 2024)",
        "ESL Hall of Fame Inductee (2019)",
        "PC Player of the Decade (2025)",
    ]
    
    for ach in achievements:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(ach)
        run.font.color.rgb = TEXT_COLOR
    
    doc.add_paragraph(
        "Faker's career spans over a decade of competitive excellence, marked by numerous records "
        "that may never be broken. He is the only player to have won the World Championship six times, "
        "the only player to defend the title twice, and holds virtually every significant record in LCK history."
    )
    
    doc.add_page_break()

    # =========================================================================
    # SECTION 2: THE RISE
    # =========================================================================
    add_styled_heading(doc, "2. The Rise: 2013-2014", 1)
    
    era_para = doc.add_paragraph()
    run = era_para.add_run("Era I: The Prodigy Emerges")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = ACCENT_COLOR
    era_sub = doc.add_paragraph("From unknown talent to World Champion in under a year")
    era_sub.runs[0].italic = True
    
    add_styled_heading(doc, "2.1 Rookie Sensation", 2)
    
    doc.add_paragraph(
        "Faker's professional debut came on February 6, 2013, when SK Telecom T1 #2 faced CJ Entus Blaze "
        "in the OGN Champions Spring 2013. What made this debut legendary was not just his performance, "
        "but the statement it made. Facing Ambition, widely considered the best mid laner in Korea at the "
        "time, Faker dominated the lane and established himself as an immediate threat."
    )
    
    # Timeline
    timeline_data = [
        ("Feb 2013", "Professional Debut", "Debuts against CJ Entus Blaze, defeating Ambition in lane"),
        ("Aug 2013", "OGN Summer Championship", 'Defeats KT Rolster Bullets in finals; the famous "Faker vs Ryu" Zed play'),
        ("Oct 2013", "First World Championship", "SKT T1 sweeps Royal Club 3-0 in the finals at Staples Center"),
    ]
    
    for date, title, desc in timeline_data:
        p = doc.add_paragraph()
        date_run = p.add_run(f"{date}: ")
        date_run.bold = True
        date_run.font.color.rgb = ACCENT_COLOR
        title_run = p.add_run(f"{title}\n")
        title_run.bold = True
        p.add_run(desc)
    
    add_styled_heading(doc, "2.2 First World Championship", 2)
    
    doc.add_paragraph(
        "The Season 3 World Championship was Faker's coronation. At just 17 years old, he led SKT T1 "
        "to a dominant 15-3 record throughout the tournament. The finals against Royal Club were "
        "particularly one-sided, with SKT winning 3-0 in what remains one of the most lopsided "
        "World Championship finals in history."
    )
    
    add_data_table(doc, 
        ["Statistic", "Value", "Tournament Rank"],
        [
            ["Kill Participation", "36%+ of team kills", "Top 3"],
            ["Gold Per Minute", "492 GPM (vs Ahri game)", "#1"],
            ["Tournament Record", "15W - 3L", "Best in tournament"],
            ["Finals Performance", "3-0 sweep", "Most dominant final"],
        ]
    )
    
    add_quote_box(doc, "He is the Lionel Messi of League of Legends.", "Maknoon, Professional Player")
    
    doc.add_page_break()

    # =========================================================================
    # SECTION 3: DYNASTY ERA
    # =========================================================================
    add_styled_heading(doc, "3. Dynasty Era: 2015-2017", 1)
    
    era_para = doc.add_paragraph()
    run = era_para.add_run("Era II: The Dynasty")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = ACCENT_COLOR
    era_sub = doc.add_paragraph("Three World Finals appearances, two more titles, and unprecedented dominance")
    era_sub.runs[0].italic = True
    
    add_styled_heading(doc, "3.1 Back-to-Back World Titles", 2)
    
    doc.add_paragraph(
        "After a rebuilding 2014, Faker returned with a vengeance in 2015. Despite sharing time with "
        "substitute mid laner Easyhoon, Faker's impact was undeniable. The team won LCK Spring, finished "
        "second at MSI (losing to EDG in a heartbreaking Game 5), and then captured both LCK Summer and "
        "the 2015 World Championship."
    )
    
    timeline_data = [
        ("2015", "Second World Championship", "SKT defeats Koo Tigers 3-1 in finals; Faker's Ryze becomes legendary"),
        ("2016", "Historic Three-Peat Begins", "Wins MSI, IEM World Championship, and Worlds 2016; named Finals MVP"),
        ("Jul 2016", "1,000 LCK Kills", "Becomes first player in LCK history to reach 1,000 career kills"),
    ]
    
    for date, title, desc in timeline_data:
        p = doc.add_paragraph()
        date_run = p.add_run(f"{date}: ")
        date_run.bold = True
        date_run.font.color.rgb = ACCENT_COLOR
        title_run = p.add_run(f"{title}\n")
        title_run.bold = True
        p.add_run(desc)
    
    doc.add_paragraph(
        "The 2016 World Championship cemented Faker's status as the greatest. In a thrilling finals "
        "against Samsung Galaxy, SKT prevailed 3-2, with Faker earning his first Finals MVP award."
    )
    
    add_styled_heading(doc, "3.2 The 2017 Heartbreak", 2)
    
    doc.add_paragraph(
        "2017 was a year of extremes. Faker carried a struggling SKT roster through MSI to victory, "
        "earning MVP honors. His Galio performances at MSI became the stuff of legend. However, the "
        "World Championship finals against Samsung Galaxy brought heartbreak—a 3-0 defeat that saw "
        "Faker reduced to tears on stage."
    )
    
    doc.add_paragraph(
        "The image of Faker crying after the 2017 Worlds finals defeat became one of the most powerful "
        'moments in esports history. Despite the loss, it humanized the "Unkillable Demon King" and '
        "showed his passion for competition."
    )
    
    add_stats_table(doc, [
        ("World Finals", "8"),
        ("MSI Titles", "2×"),
        ("Titles in 2016", "3"),
    ])
    
    doc.add_page_break()

    # =========================================================================
    # SECTION 4: DARK AGE
    # =========================================================================
    add_styled_heading(doc, "4. The Dark Age: 2018-2022", 1)
    
    era_para = doc.add_paragraph()
    run = era_para.add_run("Era III: The Struggle")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = ACCENT_COLOR
    era_sub = doc.add_paragraph("Five years without international glory, but legends never fade")
    era_sub.runs[0].italic = True
    
    add_styled_heading(doc, "4.1 Rebuilding Years", 2)
    
    doc.add_paragraph(
        "Following the 2017 World Championship defeat, SKT (later rebranded to T1) entered a period "
        "of rebuilding. The roster went through numerous changes, with Faker often the only constant. "
        "Despite winning LCK titles in 2019 and 2020, international success remained elusive."
    )
    
    timeline_data = [
        ("2018", "Roster Struggles", "SKT fails to qualify for Worlds for the first time in Faker's career"),
        ("2019", "Return to Form", "Wins LCK Spring and Summer; reaches Worlds semifinals"),
        ("Feb 2020", "Becomes Part Owner", "Signs three-year extension, becomes co-owner of T1 Entertainment"),
        ("Mar 2020", "2,000 LCK Kills", "First player to reach 2,000 career kills in LCK"),
    ]
    
    for date, title, desc in timeline_data:
        p = doc.add_paragraph()
        date_run = p.add_run(f"{date}: ")
        date_run.bold = True
        date_run.font.color.rgb = ACCENT_COLOR
        title_run = p.add_run(f"{title}\n")
        title_run.bold = True
        p.add_run(desc)
    
    add_styled_heading(doc, "4.2 Near Misses", 2)
    
    doc.add_paragraph(
        "The 2022 season brought T1 tantalizingly close to glory. The team went 18-0 in the LCK Spring "
        "Split—the first perfect regular season in LCK history—and reached the World Championship finals. "
        "However, they fell to DRX in a dramatic five-game series."
    )
    
    milestones = [
        "First player to reach 500 games in LCK (August 2019)",
        "First player to reach 2,000 kills in LCK (March 2020)",
        "First player to reach 100 international wins (October 2019)",
        "Longest tenure with a single organization in esports history",
        "Offered $20 million/year by Chinese team, chose to stay with T1",
        "First player to play 1,000 professional games (February 2022)",
    ]
    
    for m in milestones:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(m)
    
    add_quote_box(doc, 
        "I thought winning another world championship in SKT would be for the best.",
        "Faker, on rejecting the $20M offer"
    )
    
    doc.add_page_break()

    # =========================================================================
    # SECTION 5: RETURN OF THE KING
    # =========================================================================
    add_styled_heading(doc, "5. The Return of the King: 2023-2025", 1)
    
    era_para = doc.add_paragraph()
    run = era_para.add_run("Era IV: Redemption")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = ACCENT_COLOR
    era_sub = doc.add_paragraph("The unprecedented three-peat and immortal legacy")
    era_sub.runs[0].italic = True
    
    add_styled_heading(doc, "5.1 Fourth World Title", 2)
    
    doc.add_paragraph(
        "The 2023 season began with challenges—Faker suffered an arm injury in July that sidelined "
        "him for four weeks. During his absence, T1's record plummeted from 6-2 to 7-9. His return "
        "on August 2nd immediately turned the tide."
    )
    
    doc.add_paragraph(
        "The 2023 World Championship finals were held in Seoul, South Korea—Faker's home country. "
        "In front of a home crowd, he delivered a dominant 3-0 sweep against Weibo Gaming, ending "
        "six years of international drought and becoming the only player in history to win four "
        "World Championships."
    )
    
    add_styled_heading(doc, "5.2 Historic Three-Peat", 2)
    
    doc.add_paragraph(
        "What followed in 2024 and 2025 defied all expectations. Faker led T1 to back-to-back World "
        "Championships, completing the first-ever three-peat in League of Legends history."
    )
    
    add_stats_table(doc, [
        ("5th World Title", "2024"),
        ("6th World Title", "2025"),
        ("Age at 6th Title", "28"),
    ])
    
    timeline_data = [
        ("Feb 2024", "600 LCK Wins", "First player to reach 600 wins in a single region"),
        ("Feb 2024", "3,000 LCK Kills", "First player to reach 3,000 career kills in LCK"),
        ("Apr 2024", "5,000 LCK Assists", "First player to reach 5,000 career assists in LCK"),
        ("Oct 2024", "100 Worlds Wins", "First player to reach 100 wins at World Championships"),
        ("Nov 2024", "Fifth World Championship", "Defeats Bilibili Gaming; named Finals MVP (second time)"),
        ("Nov 2025", "Sixth World Championship", "Defeats KT Rolster in dramatic 5-game finals; completes three-peat"),
    ]
    
    for date, title, desc in timeline_data:
        p = doc.add_paragraph()
        date_run = p.add_run(f"{date}: ")
        date_run.bold = True
        date_run.font.color.rgb = ACCENT_COLOR
        title_run = p.add_run(f"{title}\n")
        title_run.bold = True
        p.add_run(desc)
    
    doc.add_paragraph(
        "In May 2024, Riot Games announced Faker as the inaugural inductee into the LoL Esports "
        "Hall of Legends. The honor was bestowed during the LCK broadcast, with a special ceremony "
        "celebrating his unprecedented contributions to the game and esports."
    )
    
    doc.add_page_break()

    # =========================================================================
    # SECTION 6: STATISTICAL ANALYSIS
    # =========================================================================
    add_styled_heading(doc, "6. Statistical Analysis", 1)
    
    add_styled_heading(doc, "6.1 Career Statistics", 2)
    
    doc.add_paragraph(
        "Faker's career statistics paint the picture of sustained excellence unmatched in esports history. "
        "Across more than 1,600 professional games, he has maintained elite performance metrics while "
        "adapting to countless meta changes."
    )
    
    add_stats_table(doc, [
        ("Total Games", "1,629+"),
        ("Career Record", "984 - 493"),
        ("Win Rate", "66.6%"),
        ("Career KDA", "4.1"),
        ("CS Per Minute", "8.8"),
        ("Gold Per Minute", "412"),
        ("Kill Participation", "64.2%"),
        ("Solo Kills", "340+"),
        ("Unique Champions", "89"),
    ])
    
    add_data_table(doc, 
        ["Record Category", "Value", "Status"],
        [
            ["Career Kills", "3,500+", "All-time #1"],
            ["Career Assists", "5,000+", "All-time #1"],
            ["Career Wins", "700+", "All-time #1"],
            ["Games Played", "1,000+", "All-time #1"],
            ["Titles Won", "10", "All-time #1"],
        ]
    )
    
    add_styled_heading(doc, "6.2 Champion Pool", 2)
    
    doc.add_paragraph(
        "Faker's champion pool is one of the most diverse in professional League of Legends history. "
        "He has played 89 unique champions professionally, demonstrating remarkable versatility."
    )
    
    add_data_table(doc,
        ["Rank", "Champion", "Games", "Wins", "Win Rate", "KDA"],
        [
            ["1", "Azir", "204", "143", "70.1%", "3.9"],
            ["2", "Orianna", "116", "84", "72.4%", "5.3"],
            ["3", "Ryze", "102", "65", "63.7%", "4.0"],
            ["4", "Corki", "83", "55", "66.3%", "4.3"],
            ["5", "Ahri", "79", "53", "67.1%", "4.6"],
            ["6", "Galio", "75", "46", "61.3%", "4.3"],
            ["7", "Viktor", "73", "52", "71.2%", "4.2"],
            ["8", "Taliyah", "69", "48", "69.6%", "4.6"],
            ["9", "LeBlanc", "63", "51", "81.0%", "5.7"],
            ["10", "Lissandra", "49", "39", "79.6%", "4.1"],
        ],
        highlight_first=True
    )
    
    doc.add_paragraph().add_run("Perfect and Near-Perfect Champions:").bold = True
    
    perfect_champs = [
        ("Zilean", "10 Games", "100% WR"),
        ("Diana", "4 Games", "100% WR"),
        ("Riven", "3 Games", "100% WR"),
        ("Gragas", "16 Games", "87.5% WR"),
        ("Fizz", "10 Games", "90% WR"),
        ("LeBlanc", "63 Games", "81% WR"),
    ]
    
    for champ, games, wr in perfect_champs:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f"{champ}: ")
        run.bold = True
        p.add_run(f"{games}, {wr}")
    
    doc.add_page_break()

    # =========================================================================
    # SECTION 7: LEGACY
    # =========================================================================
    add_styled_heading(doc, "7. Legacy & Impact", 1)
    
    add_styled_heading(doc, "Cultural Significance", 2)
    
    doc.add_paragraph(
        "Faker's impact extends far beyond statistics and trophies. He has become a symbol of "
        "excellence, dedication, and longevity in an industry known for player burnout and short careers. "
        "His decision to remain with T1 despite offers reportedly worth $20 million annually from Chinese "
        "organizations speaks to a loyalty that has endeared him to fans worldwide."
    )
    
    impact_points = [
        "Proven that longevity is possible in professional gaming",
        "Demonstrated that peak performance can be maintained into late twenties",
        "Elevated player compensation standards globally",
        "Become a global ambassador for Korean esports",
        "Inspired countless players to pursue professional gaming",
    ]
    
    for point in impact_points:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(point)
    
    add_styled_heading(doc, "Awards and Recognition", 2)
    
    add_data_table(doc,
        ["Tournament", "Year(s)", "Count"],
        [
            ["World Championship", "2013, 2015, 2016, 2023, 2024, 2025", "6×"],
            ["LCK Championship", "2013-2017, 2019-2020, 2022-2024", "10×"],
            ["Mid-Season Invitational", "2016, 2017", "2×"],
            ["Esports World Cup", "2024", "1×"],
            ["IEM World Championship", "2016", "1×"],
        ]
    )
    
    add_styled_heading(doc, "Most Iconic Moments", 2)
    
    iconic_moments = [
        ("Faker vs Ryu - Zed Outplay (2013)", 
         "The greatest play in League of Legends history. A perfect 1v1 outplay that became legendary."),
        ("The Tears at 2017 Worlds", 
         'After losing 0-3, cameras captured Faker crying on stage, humanizing the "Unkillable Demon King."'),
        ("The 2023 Homecoming", 
         "Winning Worlds in Seoul after six years without an international title."),
        ("The Three-Peat Completion (2025)", 
         "Defeating KT Rolster in a dramatic 5-game finals to complete the first-ever Worlds three-peat."),
    ]
    
    for title, desc in iconic_moments:
        p = doc.add_paragraph()
        run = p.add_run(f"{title}: ")
        run.bold = True
        p.add_run(desc)
    
    doc.add_page_break()

    # =========================================================================
    # SECTION 8: CONCLUSION
    # =========================================================================
    add_styled_heading(doc, "8. Conclusion", 1)
    
    doc.add_paragraph(
        "The story of Faker is the story of esports itself. From a talented teenager making his debut "
        "in 2013 to a 29-year-old completing an unprecedented three-peat in 2025, Lee \"Faker\" Sang-hyeok "
        "has defined what it means to be a professional gamer."
    )
    
    doc.add_paragraph(
        "His six World Championship titles stand alone in the record books. His 10 LCK titles represent "
        "domestic dominance unmatched in any major region. His 12+ years with a single organization "
        "demonstrate loyalty that transcends the transactional nature of modern sports."
    )
    
    add_stats_table(doc, [
        ("World Championships", "6"),
        ("LCK Championships", "10"),
        ("World Finals", "8"),
        ("Games Played", "1,600+"),
        ("Career Win Rate", "66.6%"),
        ("Years with T1", "12+"),
    ])
    
    doc.add_paragraph(
        "As the inaugural inductee into the LoL Esports Hall of Legends, Faker has already secured "
        "his place in history. Yet remarkably, his story continues. Whether he adds more titles to "
        "his collection or mentors the next generation of players, his impact on League of Legends "
        "and esports as a whole is immeasurable."
    )
    
    add_quote_box(doc, "There is only one Faker. There will only ever be one Faker.", "League of Legends Community")
    
    # Final statement
    for _ in range(2):
        doc.add_paragraph()
    
    final = doc.add_paragraph()
    final.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run1 = final.add_run("The Unkillable Demon King.\n")
    run1.bold = True
    run1.font.size = Pt(16)
    run1.font.color.rgb = ACCENT_COLOR
    run2 = final.add_run("The Greatest of All Time.")
    run2.bold = True
    run2.font.size = Pt(16)
    run2.font.color.rgb = ACCENT_COLOR
    
    doc.add_page_break()

    # =========================================================================
    # REFERENCES
    # =========================================================================
    add_styled_heading(doc, "References", 1)
    
    references = [
        'LoL Esports (2024). "LoL Esports Welcomes Faker to Hall of Legends." LoL Esports Official.',
        'Leaguepedia (2025). "Faker - Player Profile and Statistics." Leaguepedia Wiki.',
        'Liquipedia (2025). "Faker Biography and Achievements." Liquipedia League of Legends.',
        'Games of Legends (2025). "Faker Career Statistics." gol.gg.',
        'Wikipedia (2025). "Faker (gamer)." Wikipedia.',
        'Faker (2016). "Unkillable." The Players\' Tribune.',
        'Inven Global (2023). "A Decade of Faker: Ranking the ten years of the T1 mid laner\'s career." Inven Global.',
        'Red Bull (2021). "10 things you didn\'t know about League of Legends supremo Faker." Red Bull Gaming.',
        'ESPN Esports (2016). "A global phenomenon: SK Telecom T1 Faker, esports\' Michael Jordan." ESPN.',
        'Dot Esports (2024). "Faker goes above and beyond as T1 defend LoL Worlds title." Dot Esports.',
    ]
    
    for i, ref in enumerate(references, 1):
        p = doc.add_paragraph()
        run = p.add_run(f"[{i}] ")
        run.bold = True
        p.add_run(ref)
        p.paragraph_format.left_indent = Cm(1)
        p.paragraph_format.first_line_indent = Cm(-1)
    
    # Save document
    output_path = "/Users/vietth18.te/thvroyal/kimi-skills/faker-career-analysis.docx"
    doc.save(output_path)
    print(f"Document saved to: {output_path}")
    return output_path

if __name__ == "__main__":
    create_faker_document()
